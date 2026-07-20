import os
from pathlib import Path
from docx import Document
import cloudconvert

from utilidades import convertir_numero_a_palabras, formatear_fecha_larga

API_KEY = os.environ.get("CLOUDCONVERT_API_KEY")
if API_KEY:
    cloudconvert.configure(api_key=API_KEY)

def convertir_docx_a_pdf_cloudconvert(ruta_docx, ruta_pdf):
    if not API_KEY:
        raise ValueError("La variable CLOUDCONVERT_API_KEY no está configurada en Render.")

    job = cloudconvert.Job.create(payload={
        "tasks": {
            "importar-docx": {"operation": "import/upload"},
            "convertir-a-pdf": {
                "operation": "convert",
                "input": "importar-docx",
                "output_format": "pdf",
                "engine": "office"
            },
            "exportar-pdf": {"operation": "export/url", "input": "convertir-a-pdf"}
        }
    })

    upload_task = filter(lambda task: task['operation'] == 'import/upload', job['tasks'])
    upload_task = list(upload_task)[0]

    cloudconvert.Task.upload(file_name=ruta_docx, task=upload_task)
    completed_job = cloudconvert.Job.wait(id=job["id"])

    export_task = filter(lambda task: task['operation'] == 'export/url', completed_job['tasks'])
    export_task = list(export_task)[0]

    if export_task["status"] == "finished" and export_task["result"]["files"]:
        archivo_salida = export_task["result"]["files"][0]
        cloudconvert.download(url=archivo_salida["url"], filename=ruta_pdf)
        return ruta_pdf
    else:
        raise Exception("La conversión en CloudConvert falló en los servidores externos.")

def generar_carta_cesantias(datos_formulario):
    base_dir = Path(__file__).resolve().parent
    ruta_plantilla = base_dir / "plantilla.docx"

    nombre_base = f"Carta_Retiro_{datos_formulario.get('cedula', 'temporal')}"
    ruta_docx_salida = str(base_dir / f"{nombre_base}.docx")
    ruta_pdf_salida = str(base_dir / f"{nombre_base}.pdf")

    doc = Document(ruta_plantilla)

    # Preparar los datos ingresados y limpiar textos
    nombre = datos_formulario.get('nombre', '')
    cedula = datos_formulario.get('cedula', '')
    valor_num = datos_formulario.get('valor', '0')
    fecha_maxima = datos_formulario.get('fecha_maxima', '')
    fecha_expedicion = datos_formulario.get('fecha_expedicion', '')

    # Convertir el valor numérico a letras de forma automática si es un número válido
    try:
        # Quitamos puntos o comas que el usuario haya puesto para que la función no falle
        valor_limpio = int(valor_num.replace('.', '').replace(',', '').strip())
        valor_letras = convertir_numero_a_palabras(valor_limpio).upper() + " PESOS M/CTE"
    except Exception:
        valor_letras = ""

    # Diccionario con todos los reemplazos (en Mayúsculas como en tu Word)
    reemplazos = {
        "{{NOMBRE}}": nombre,
        "{{CEDULA}}": cedula,
        "{{VALOR}}": valor_num,
        "{{VALOR_LETRAS}}": valor_letras,
        "{{FECHA_MAXIMA}}": fecha_maxima,
        "{{FECHA_EXPEDICION}}": fecha_expedicion
    }

    # 1. Reemplazar datos en los párrafos normales del documento
    for p in doc.paragraphs:
        for clave, valor in reemplazos.items():
            if clave in p.text:
                p.text = p.text.replace(clave, valor)

    # 2. Reemplazar datos dentro de las TABLAS (esencial para tu formato)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    for clave, valor in reemplazos.items():
                        if clave in p.text:
                            p.text = p.text.replace(clave, valor)

    doc.save(ruta_docx_salida)

    try:
        convertir_docx_a_pdf_cloudconvert(ruta_docx_salida, ruta_pdf_salida)
        return ruta_docx_salida, ruta_pdf_salida
    except Exception as e:
        print(f"Error convirtiendo a PDF: {e}")
        return ruta_docx_salida, None
